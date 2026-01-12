"""Word document (.docx) parser."""

import io
import logging

from .base import BaseParser, ContentBlock, ContentType, ParsedDocument, TableContent

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Parser for Word documents (.docx).

    Uses python-docx library to extract text, tables, and metadata.
    """

    @property
    def supported_mime_types(self) -> list[str]:
        """Return list of supported MIME types."""
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]

    async def parse(
        self, content: bytes, metadata: dict | None = None,
    ) -> ParsedDocument:
        """Parse Word document.

        Args:
            content: Raw .docx file bytes.
            metadata: Optional metadata to include.

        Returns:
            ParsedDocument with extracted content.

        Raises:
            ValueError: If document cannot be parsed.
        """
        from docx import Document

        metadata = metadata or {}

        try:
            doc = Document(io.BytesIO(content))
        except Exception as e:
            raise ValueError(f"Failed to parse Word document: {e}") from e

        blocks: list[ContentBlock] = []
        tables: list[TableContent] = []
        full_text: list[str] = []
        position = 0

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
                blocks.append(
                    ContentBlock(
                        content_type=ContentType.TEXT,
                        content=text,
                        position=position,
                        metadata={
                            "style": para.style.name if para.style else None,
                        },
                    ),
                )
                position += 1

        # Extract tables
        for table in doc.tables:
            try:
                table_content = self._extract_table(table)
                if table_content:
                    tables.append(table_content)
                    # Add table as content block
                    blocks.append(
                        ContentBlock(
                            content_type=ContentType.TABLE,
                            content=self._table_to_text(table_content),
                            position=position,
                        ),
                    )
                    position += 1
            except Exception as e:
                logger.warning(f"Failed to extract table: {e}")

        # Extract document properties
        core_props = doc.core_properties
        created_date = None
        modified_date = None

        if core_props.created:
            created_date = str(core_props.created)
        if core_props.modified:
            modified_date = str(core_props.modified)

        return ParsedDocument(
            text="\n\n".join(full_text),
            blocks=blocks,
            tables=tables,
            title=core_props.title or None,
            author=core_props.author or None,
            created_date=created_date,
            modified_date=modified_date,
            metadata=metadata,
        )

    def _extract_table(self, table) -> TableContent | None:
        """Extract table content from a docx table.

        Args:
            table: python-docx Table object.

        Returns:
            TableContent or None if extraction fails.
        """
        rows_data: list[list[str]] = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Get cell text, handling merged cells
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            rows_data.append(row_data)

        if not rows_data:
            return None

        # Use first row as headers
        headers = rows_data[0]
        rows = rows_data[1:] if len(rows_data) > 1 else []

        return TableContent(headers=headers, rows=rows)

    def _table_to_text(self, table: TableContent) -> str:
        """Convert TableContent to plain text.

        Args:
            table: TableContent to convert.

        Returns:
            Plain text representation of the table.
        """
        lines = []

        # Headers
        lines.append(" | ".join(table.headers))
        lines.append("-" * len(lines[-1]))

        # Rows
        for row in table.rows:
            lines.append(" | ".join(row))

        return "\n".join(lines)
